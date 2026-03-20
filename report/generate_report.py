#!/usr/bin/env python3
"""AI自動法廷システム 実験レポート Word生成スクリプト

システム概要・技術仕様・実行結果・課題・展望をまとめた
研究・ゼミ発表用レポートを .docx 形式で出力する。
"""

import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUTS_DIR = BASE_DIR / "outputs_agent"
RUN_DIR = OUTPUTS_DIR / "run_20260320_154336"  # メインの結果
OUTPUT_FILE = Path(__file__).resolve().parent / "experiment_report.docx"


def load_json(path: Path):
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


def set_cell_shading(cell, color: str):
    """セルの背景色を設定"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color, qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """スタイル付きテーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # ヘッダー
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")

    # データ行
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val) if val else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    # 列幅
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    return table


def truncate(text: str, max_len: int = 200) -> str:
    if not text:
        return ""
    return text[:max_len] + "…" if len(text) > max_len else text


def main():
    doc = Document()

    # ── デフォルトフォント設定 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "游明朝"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

    # ── タイトル ──
    title = doc.add_heading("AI自動法廷シミュレーションシステム\n実験レポート", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}").font.size = Pt(10)

    # ================================================================
    # 1. システム概要
    # ================================================================
    doc.add_heading("1. システム概要", level=1)

    doc.add_heading("1.1 目的", level=2)
    doc.add_paragraph(
        "本システムは、大規模言語モデル（LLM）を活用して日本の民事訴訟手続をシミュレーションする"
        "AI自動法廷システムである。原告代理人・被告代理人・書記官・裁判官の各ロールをLLMが担い、"
        "争点整理から判決に至る一連の訴訟手続を自動的に実行する。"
    )

    doc.add_heading("1.2 アーキテクチャ", level=2)
    doc.add_paragraph(
        "本システムは「ハイブリッド型」アーキテクチャを採用している。全体のフロー制御は"
        "パイプライン型ワークフロー（Python）が担い、各ロール内部ではOpenAI Agents SDKを用いた"
        "エージェントがツール呼び出しを含む自律的な推論を行う。"
    )
    doc.add_paragraph(
        "この設計は、民事訴訟手続のフェーズが法的に固定されている（訴状→答弁→争点整理→弁論→"
        "合議→判決）ことを反映し、フロー全体の予測可能性・再現性を維持しつつ、"
        "各フェーズ内での法律知識検索の柔軟性を確保するものである。"
    )

    # フロー図
    doc.add_heading("1.3 処理フロー", level=2)
    flow_text = (
        "┌─────────────────────────────────────────────────┐\n"
        "│  入力（事案テキスト＋候補条文）                  │\n"
        "└─────────────┬───────────────────────────────────┘\n"
        "              ▼\n"
        "┌─────────────────────────────────────────────────┐\n"
        "│  書記官：事案整理・争点抽出                      │\n"
        "└─────────────┬───────────────────────────────────┘\n"
        "              ▼\n"
        "┌─────────────────────────────────────────────────┐\n"
        "│  ┌─ 原告代理人：主張提出  ─┐                    │\n"
        "│  │  被告代理人：反論提出    │ × N ラウンド       │\n"
        "│  │  書記官：争点表更新      │                    │\n"
        "│  │  裁判官：継続判定        │                    │\n"
        "│  └─────────────────────────┘                    │\n"
        "└─────────────┬───────────────────────────────────┘\n"
        "              ▼\n"
        "┌─────────────────────────────────────────────────┐\n"
        "│  合議（陪席裁判官×2 → 裁判長）                  │\n"
        "└─────────────┬───────────────────────────────────┘\n"
        "              ▼\n"
        "┌─────────────────────────────────────────────────┐\n"
        "│  判決起案 → 批評 → 最終判決                     │\n"
        "└─────────────────────────────────────────────────┘\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(flow_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    # ================================================================
    # 2. 技術仕様
    # ================================================================
    doc.add_heading("2. 技術仕様", level=1)

    doc.add_heading("2.1 使用技術", level=2)
    tech_items = [
        ("LLMフレームワーク", "OpenAI Agents SDK (openai-agents)"),
        ("言語モデル", "GPT-5"),
        ("データ検証", "Pydantic v2（strict JSON schema）"),
        ("実装言語", "Python 3.12"),
        ("出力形式", "JSON（フェーズごとのスナップショット）"),
    ]
    add_styled_table(
        doc,
        ["項目", "詳細"],
        tech_items,
        col_widths=[5, 12],
    )

    doc.add_heading("2.2 ロール一覧とツール割り当て", level=2)
    doc.add_paragraph(
        "各ロールはOpenAI Agents SDKのAgentとして生成され、"
        "ロールに応じたツール（関数呼び出し）が割り当てられる。"
    )
    role_rows = [
        ("書記官 (clerk)", "事案テキストの構造化・争点表の更新", "なし"),
        ("原告代理人 (plaintiff)", "争点ごとの法的主張の構成", "条文検索, 判例検索"),
        ("被告代理人 (defendant)", "原告主張への反論の構成", "条文検索, 判例検索"),
        ("裁判官 (judge)", "弁論継続の要否判定", "なし"),
        ("陪席裁判官 (associate)", "争点ごとの法的意見形成", "条文検索"),
        ("裁判長 (presiding)", "合議取りまとめ・判決起案", "条文検索"),
        ("批評者 (critic)", "判決草案の問題点指摘", "なし"),
    ]
    add_styled_table(
        doc,
        ["ロール", "役割", "利用可能ツール"],
        role_rows,
        col_widths=[5, 8, 4],
    )

    doc.add_heading("2.3 ツール詳細", level=2)
    doc.add_paragraph(
        "search_statutes（条文検索）: 候補条文リスト内をキーワードで全文検索し、"
        "関連条文を返す。現時点では入力時に指定された候補条文のみが検索対象である。",
        style="List Bullet",
    )
    doc.add_paragraph(
        "search_case_law（判例検索）: 現在はプレースホルダ実装であり、"
        "「判例検索機能は準備中です」というメッセージを返す。今後の拡張ポイントである。",
        style="List Bullet",
    )

    # ================================================================
    # 3. 実験設定
    # ================================================================
    doc.add_heading("3. 実験設定", level=1)

    doc.add_heading("3.1 入力事例", level=2)
    case_record = load_json(RUN_DIR / "case_record.json") or {}
    doc.add_paragraph(f"事件名: {case_record.get('title', '')}")
    doc.add_paragraph(case_record.get("raw_case_text", "").split("\n\n事案：\n")[-1])

    doc.add_heading("3.2 候補条文", level=2)
    statutes = case_record.get("candidate_statutes", [])
    statute_rows = [(s["citation"], s["text"]) for s in statutes]
    add_styled_table(
        doc,
        ["条文", "内容"],
        statute_rows,
        col_widths=[4, 13],
    )

    doc.add_heading("3.3 実行パラメータ", level=2)
    meta = load_json(RUN_DIR / "meta.json") or {}
    param_rows = [
        ("モデル", "gpt-5"),
        ("最大ラウンド数", str(meta.get("max_rounds", 2))),
        ("実際ラウンド数", str(meta.get("current_round", 1))),
        ("終了理由", "裁判官が新争点なしと判断"),
    ]
    add_styled_table(doc, ["パラメータ", "値"], param_rows, col_widths=[5, 12])

    # ================================================================
    # 4. 実行結果
    # ================================================================
    doc.add_heading("4. 実行結果", level=1)

    # 4.1 争点整理
    doc.add_heading("4.1 争点整理", level=2)
    issues = load_json(RUN_DIR / "issue_table_current.json") or []
    doc.add_paragraph(f"書記官により{len(issues)}件の争点が抽出された。")
    for iss in issues:
        doc.add_paragraph(
            f"{iss['issue_id']}: {iss['title']}",
            style="List Bullet",
        )

    # 4.2 原告・被告の主張対照
    doc.add_heading("4.2 当事者主張の対照", level=2)
    for iss in issues:
        doc.add_heading(f"{iss['issue_id']}: {iss['title']}", level=3)
        claim_rows = [
            ("原告（丙）の主張", truncate(iss.get("plaintiff_argument", ""), 500)),
            ("被告（乙）の主張", truncate(iss.get("defendant_argument", ""), 500)),
        ]
        add_styled_table(doc, ["当事者", "主張要旨"], claim_rows, col_widths=[4, 13])
        doc.add_paragraph()  # spacing

    # 4.3 合議結果
    doc.add_heading("4.3 合議結果", level=2)
    deliberations = load_json(RUN_DIR / "deliberations.json") or []
    judge_names = {
        "associate_judge_1": "陪席裁判官1",
        "associate_judge_2": "陪席裁判官2",
        "presiding_judge": "裁判長",
    }
    if deliberations:
        opinions = deliberations[0].get("opinions", [])
        delib_rows = []
        for op in opinions:
            delib_rows.append((
                judge_names.get(op["speaker"], op["speaker"]),
                op.get("issue_id", ""),
                op.get("decision", ""),
                ", ".join(op.get("related_statutes", [])),
            ))
        add_styled_table(
            doc,
            ["裁判官", "争点", "判断", "適用条文"],
            delib_rows,
            col_widths=[4, 3, 5, 5],
        )

    doc.add_paragraph()
    doc.add_paragraph("3名の裁判官全員が、両争点について同一の結論に達した。")

    # 4.4 判決結果
    doc.add_heading("4.4 最終判決", level=2)
    final = load_json(RUN_DIR / "final_judgment.json") or {}
    doc.add_paragraph(f"主文: {final.get('conclusion', '')}")
    doc.add_paragraph()
    for fi in final.get("issues", []):
        doc.add_heading(f"{fi['issue_id']}: {fi.get('issue_title', '')}", level=3)
        doc.add_paragraph(f"判断: {fi.get('decision', '')}")
        doc.add_paragraph(fi.get("reasoning", ""))

    # 4.5 批評と対応
    doc.add_heading("4.5 判決批評と最終判決での改善", level=2)
    critique = load_json(RUN_DIR / "critique_log.json") or {}
    criticisms = critique.get("criticisms", [])
    doc.add_paragraph(
        f"判決草案に対し、批評エージェントが{len(criticisms)}件の問題点を指摘した。"
        "最終判決ではこれらの批評を踏まえた改訂が行われている。"
    )

    # 全最終判決reasoning結合
    full_final = " ".join(i.get("reasoning", "") for i in final.get("issues", []))

    crit_rows = []
    for i, c in enumerate(criticisms, 1):
        problem = c.get("problem", "")
        # 簡易対応チェック
        keywords = {
            "借地借家法": "借地借家法",
            "賃貸人地位承継": "賃貸人地位",
            "善意・悪意": "善意",
            "176条・177条": "176条",
            "前提事実": "前提事実",
            "信義則": "信義則",
            "解除権": "解除",
            "請求範囲": "部分",
        }
        addressed = "要確認"
        for key, term in keywords.items():
            if key in problem and term in full_final:
                addressed = "対応済"
                break
        crit_rows.append((
            str(i),
            c.get("target_issue_id", ""),
            problem,
            addressed,
        ))

    add_styled_table(
        doc,
        ["No.", "対象争点", "指摘事項", "対応状況"],
        crit_rows,
        col_widths=[1.5, 3, 10, 3],
    )

    # ================================================================
    # 5. 現状の限界と課題
    # ================================================================
    doc.add_heading("5. 現状の限界と課題", level=1)

    limitations = [
        (
            "判例検索ツールの未実装",
            "現在、判例検索ツール（search_case_law）はプレースホルダのみであり、"
            "実際の判例データベースへの接続は行われていない。法的推論の質を高めるには、"
            "裁判所判例検索システム等との連携が不可欠である。"
        ),
        (
            "条文検索の限定性",
            "条文検索は入力時に指定された候補条文リスト内のキーワード検索に留まり、"
            "六法全書や特別法の全文検索には対応していない。"
            "本実験でも、批評で指摘された借地借家法10条は候補条文に含まれていなかったが、"
            "最終判決では言及されており、LLMの事前学習知識に依存している。"
        ),
        (
            "事実認定の深度",
            "本実験では全事実が「争いなし（undisputed）」として処理されており、"
            "事実認定の争いがある事例（証拠の信用性評価等）への対応は未検証である。"
        ),
        (
            "信義則・権利濫用の評価",
            "法的原則（信義則・権利濫用）の適用判断は高度に文脈依存であり、"
            "LLMによる判断の妥当性検証が困難である。"
            "本実験でも、批評で指摘された信義則の検討について、"
            "最終判決では否定の結論を示しているが、その論証の十分性は更なる検討を要する。"
        ),
        (
            "コストと実行時間",
            "Agents SDKのツール呼び出しループにより、1回の実行あたりのAPIコール数が"
            "パイプライン型と比べて増加する。大量の事例を処理する場合のコスト最適化が課題である。"
        ),
    ]
    for title_text, desc in limitations:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(desc)

    # ================================================================
    # 6. 今後の展望
    # ================================================================
    doc.add_heading("6. 今後の展望", level=1)

    prospects = [
        (
            "判例データベースの接続",
            "裁判所判例検索システムや法律データベース（TKCローライブラリ等）と連携し、"
            "エージェントが関連判例を自律的に検索・引用できるようにする。"
            "これにより、法的推論の根拠がLLMの事前学習知識からデータベースに基づく"
            "検証可能な情報へと移行する。"
        ),
        (
            "複数事例での精度評価",
            "様々な法分野（契約法・不法行為法・物権法等）の事例を用いて"
            "システムの判断精度を横断的に評価する。"
            "判例集の解説と比較することで、AIの法的推論能力の定量的な測定を目指す。"
        ),
        (
            "事実認定争いのある事例への拡張",
            "証拠の信用性評価や経験則の適用を含む事実認定機能を実装し、"
            "より現実的な訴訟シミュレーションを実現する。"
        ),
        (
            "ガードレールの実装",
            "法的ハルシネーション（存在しない条文・判例の引用）を検出する機構や、"
            "論理的整合性チェック等のガードレールを追加し、出力の信頼性を向上させる。"
        ),
    ]
    for title_text, desc in prospects:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(desc)

    # ── 保存 ──
    doc.save(str(OUTPUT_FILE))
    print(f"Report saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
